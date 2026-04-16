"""Parser para documentos Word (DOCX) usando python-docx."""

import io
from typing import Any

from docx import Document
from structlog import get_logger

logger = get_logger(__name__)


class WordParser:
    """Parser para extrair texto e metadados de arquivos Word (DOCX)."""

    def __init__(self) -> None:
        """Inicializa o parser Word."""

    async def extract_text(self, file_content: bytes) -> str:
        """
        Extrai texto de parágrafos e tabelas do documento Word.

        Args:
            file_content: Conteúdo binário do arquivo DOCX.

        Returns:
            Texto extraído de parágrafos e tabelas.
            Retorna string vazia em caso de erro.
        """
        if not self._validate_docx_bytes(file_content):
            logger.warning("docx_invalid_bytes", size=len(file_content))
            return ""

        try:
            doc = Document(io.BytesIO(file_content))
            extracted_parts = []

            # Extrair texto de parágrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    extracted_parts.append(text)

            # Extrair texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        extracted_parts.append(row_text)

            result = "\n\n".join(extracted_parts)

            logger.info(
                "docx_text_extracted",
                paragraph_count=len(doc.paragraphs),
                table_count=len(doc.tables),
                char_count=len(result),
            )

            return result

        except Exception as e:
            logger.error("docx_extraction_failed", error=str(e))
            return ""

    async def extract_metadata(self, file_content: bytes) -> dict[str, Any]:
        """
        Extrai metadados do documento Word.

        Args:
            file_content: Conteúdo binário do arquivo DOCX.

        Returns:
            Dicionário com metadados: paragraph_count, table_count, title, author, etc.
        """
        if not self._validate_docx_bytes(file_content):
            return {}

        metadata: dict[str, Any] = {}

        try:
            doc = Document(io.BytesIO(file_content))

            # Contagem de elementos
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            metadata["section_count"] = len(doc.sections)

            # Propriedades do documento (core properties)
            core_props = doc.core_properties

            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.category:
                metadata["category"] = core_props.category
            if core_props.created:
                # created pode ser datetime ou string
                if hasattr(core_props.created, "isoformat"):
                    metadata["created"] = core_props.created.isoformat()
                else:
                    metadata["created"] = str(core_props.created)
            if core_props.modified:
                # modified pode ser datetime ou string
                if hasattr(core_props.modified, "isoformat"):
                    metadata["modified"] = core_props.modified.isoformat()
                else:
                    metadata["modified"] = str(core_props.modified)
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.revision:
                metadata["revision"] = core_props.revision
            if core_props.version:
                metadata["version"] = core_props.version

            logger.info(
                "docx_metadata_extracted",
                paragraph_count=metadata.get("paragraph_count"),
                table_count=metadata.get("table_count"),
            )

        except Exception as e:
            logger.error("docx_metadata_extraction_failed", error=str(e))
            return {}

        return metadata

    def validate(self, file_content: bytes) -> bool:
        """
        Valida se o conteúdo é um DOCX válido.

        Args:
            file_content: Conteúdo binário do arquivo.

        Returns:
            True se for um DOCX válido, False caso contrário.
        """
        return self._validate_docx_bytes(file_content)

    def _validate_docx_bytes(self, file_content: bytes) -> bool:
        """
        Valida bytes como DOCX verificando a assinatura ZIP.

        Args:
            file_content: Conteúdo binário a validar.

        Returns:
            True se tiver assinatura ZIP válida (DOCX é um ZIP).
        """
        if not file_content or len(file_content) < 4:
            return False

        # DOCX é um ZIP file - magic number: PK (0x504B)
        return file_content[:2] == b"PK"
